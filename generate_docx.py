#!/usr/bin/env python3
"""Generate DOCX proposal from markdown and figures."""

import os
import re
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BASEDIR = os.path.dirname(__file__)
FIGDIR = os.path.join(BASEDIR, 'docs', 'figures')
OUTPATH = os.path.join(BASEDIR, 'docs', 'proposal-corning-kr.docx')

def set_cell_text(cell, text, bold=False, size=10):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, size=10)
    # Rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            set_cell_text(table.rows[r_idx + 1].cells[c_idx], val, size=10)
    return table

def create_docx():
    doc = Document()
    
    # --- Page margins ---
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # --- Default font ---
    style = doc.styles['Normal']
    font = style.font
    font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5
    
    # =========================================================================
    # TITLE PAGE
    # =========================================================================
    for _ in range(6):
        doc.add_paragraph('')
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Sub-5nm EUV 리소그래피용 정밀유리 기판\nMicro-Crack Lifecycle 관리 시뮬레이터 개발')
    run.font.size = Pt(16)
    run.bold = True
    
    doc.add_paragraph('')
    
    org = doc.add_paragraph()
    org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = org.add_run('제안 기관: 성균관대학교 SPMDL (권석준 교수 연구실)\n협력 대상: Corning Incorporated — 천안 사업장')
    run.font.size = Pt(12)
    
    doc.add_paragraph('')
    
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run('제안 일자: 2026년 2월')
    run.font.size = Pt(12)
    
    doc.add_page_break()
    
    # =========================================================================
    # SECTION 1
    # =========================================================================
    h1 = doc.add_heading('1. 연구 배경 및 필요성', level=1)
    
    h2 = doc.add_heading('EUV/High-NA 전환에 따른 substrate 정밀도 요구 급증', level=2)
    doc.add_paragraph(
        '차세대 반도체 공정이 Sub-5nm 노드로 진입하면서, EUV 리소그래피 광학계 및 photomask substrate에 '
        '요구되는 정밀도가 급격히 높아지고 있습니다. ASML EXE:5200 (High-NA, 0.55 NA) 도입과 함께 '
        'overlay 오차 허용 범위가 1.5 nm RMS 이하로 축소되었으며, 이는 glass substrate의 열적·기계적 '
        '안정성에 대한 전례 없는 수준의 요구입니다.'
    )
    
    h2 = doc.add_heading('Micro-crack: Overlay 오차와 yield 저하의 핵심 원인', level=2)
    doc.add_paragraph(
        'Glass substrate 내부의 micro-crack은 다음 경로를 통해 공정 성능을 저하시킵니다:'
    )
    bullets = [
        '국소적 CTE(열팽창계수) 이상 → 열 변형 → overlay registration 오차',
        '국소적 굴절률 변화 (Δn) → 위상 오차 → pattern placement 오차',
        'Subcritical crack growth (SCG) → 사용 중 균열 진전 → 점진적 성능 열화',
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')
    
    doc.add_paragraph(
        '현재 업계에서는 사후 검사(post-mortem inspection) 위주의 품질 관리가 이루어지고 있어, '
        '균열의 발생 시점과 진전 경로를 예측하기 어렵습니다. 이는 예기치 않은 공정 중단과 yield 손실로 이어집니다.'
    )
    
    h2 = doc.add_heading('Corning ULE 7973의 기술적 과제', level=2)
    doc.add_paragraph(
        'Corning ULE 7973은 EUV photomask substrate의 사실상 표준(de facto standard)이나, '
        '비정질(amorphous) 구조 특성상 결정질 소재 대비 균열 거동 예측이 본질적으로 어렵습니다. '
        '결정 입계(grain boundary)가 없어 crack deflection이 일어나지 않으며, '
        'subcritical crack growth 거동이 환경 조건(습도, 온도)에 민감합니다.'
    )
    
    # =========================================================================
    # SECTION 2
    # =========================================================================
    doc.add_heading('2. 연구 목표', level=1)
    doc.add_paragraph(
        '본 과제는 Glass Micro-Crack Lifecycle Simulator를 개발하여 다음 세 가지 핵심 목표를 달성하고자 합니다:'
    )
    goals = [
        'Full Lifecycle 시뮬레이터 개발: Nucleation → Propagation → Inspection → Diagnostics → Attribution의 5단계를 통합하는 physics-informed 시뮬레이션 플랫폼 구축',
        '경쟁 소재 대비 ULE 7973 우위 정량화: Zerodur, Clearceram-Z, AGC AZ, Shin-Etsu Quartz 등 5개 소재와의 체계적 비교를 통해 ULE 7973의 기술적 moat를 수치적으로 입증',
        'ML 기반 실시간 진단 시스템 프로토타입: Bayesian inference 기반의 역추론 엔진을 통해 검사 데이터로부터 균열 상태를 실시간 진단하는 프로토타입 시스템 개발',
    ]
    for i, g in enumerate(goals, 1):
        doc.add_paragraph(f'{i}. {g}')
    
    # =========================================================================
    # SECTION 3
    # =========================================================================
    doc.add_heading('3. 핵심 기술 및 방법론', level=1)
    doc.add_paragraph(
        '본 시뮬레이터는 5개의 독립 모듈로 구성되며, 각 모듈은 물리 법칙에 기반한 forward model과 '
        '데이터 기반 inverse model을 결합합니다.'
    )
    
    # Fig 0
    fig0_path = os.path.join(FIGDIR, 'fig0_schematic.png')
    if os.path.exists(fig0_path):
        doc.add_picture(fig0_path, width=Cm(15))
        cap = doc.add_paragraph('Figure 0. 시스템 구조도 — Glass Micro-Crack Lifecycle Simulator')
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].italic = True
        cap.runs[0].font.size = Pt(10)
    
    # Module table
    add_table(doc,
        ['모듈', '기능', '핵심 물리/알고리즘'],
        [
            ['M1: Nucleation Engine', '결함 분포 생성 및 nucleation 확률 계산', 'Griffith criterion, Monte Carlo, Poisson/Neyman-Scott 결함 분포'],
            ['M2: Propagation Engine', 'Subcritical crack growth 및 phase-field 균열 진전', 'Charles-Hillig velocity law, Paris fatigue law, Phase-field fracture'],
            ['M3: Inspection Forward Model', '6종 검사 기법의 신호 시뮬레이션', 'Lamb wave, Rayleigh/Mie scattering, Raman, 193nm interferometry, EELS, KFM'],
            ['M4: Inverse ML Diagnostics', '검사 신호로부터 균열 상태 역추론', 'Gaussian Process, Bayesian inference, Physics-informed features'],
            ['M5: Process Attribution', '공정 분산 분해 및 교체 시점 최적화', 'Variance decomposition, Bayesian changepoint, Cost optimization'],
        ]
    )
    
    doc.add_paragraph('')
    doc.add_paragraph(
        '현재 M1–M5 전체 모듈 구현이 완료되었으며 (156/156 unit tests 통과), '
        '시뮬레이션 기반 검증을 진행 중입니다.'
    )
    
    # =========================================================================
    # SECTION 4
    # =========================================================================
    doc.add_heading('4. 기대 효과', level=1)
    doc.add_paragraph(
        '시뮬레이션 기반 추정에 따르면, 본 시스템의 적용을 통해 다음과 같은 효과를 기대할 수 있습니다:'
    )
    
    add_table(doc,
        ['항목', '추정 효과', '근거'],
        [
            ['Overlay 오차 기여분 감소', '15–30% (mask degradation 성분)', 'M5 variance decomposition 기반'],
            ['비계획 정지 감소', '20–40%', 'Bayesian changepoint에 의한 조기 경보'],
            ['Substrate 교체 비용 절감', '연간 $100K–$500K (라인 당)', 'M5 replacement optimization 기반'],
            ['공정 최적화 시간 단축', '30–50%', 'M4 역추론에 의한 root cause 자동 진단'],
        ]
    )
    
    doc.add_paragraph('')
    doc.add_paragraph(
        '※ 상기 수치는 시뮬레이션 기반 추정치이며, 실제 효과는 공정 조건과 운영 환경에 따라 달라질 수 있습니다.',
    )
    
    doc.add_paragraph('Corning 관점의 전략적 가치:', style='List Bullet')
    strat = [
        'ULE 7973의 SCG exponent (n=20), 높은 활성화 에너지 (80 kJ/mol) 등 경쟁 우위를 정량적으로 입증하는 도구 확보',
        '고객사(TSMC, Samsung, Intel 등)에 대한 기술 차별화 자료 제공',
        'High-NA EUV 전환 시 Extreme-ULE 개발 방향에 대한 시뮬레이션 기반 가이드라인 제시',
    ]
    for s in strat:
        doc.add_paragraph(s, style='List Bullet 2')
    
    # =========================================================================
    # SECTION 5
    # =========================================================================
    doc.add_heading('5. 연구진 및 수행 체계', level=1)
    
    doc.add_heading('연구 책임자', level=2)
    doc.add_paragraph('권석준 교수 — 성균관대학교 SPMDL (Semiconductor Physics & Materials Design Lab)')
    
    doc.add_heading('수행 일정', level=2)
    add_table(doc,
        ['Phase', '기간', '주요 내용'],
        [
            ['Phase 1', '8개월', '시뮬레이터 고도화 및 Corning 실험 데이터 연동 (M1–M3 calibration)'],
            ['Phase 2', '10개월', 'ML 진단 시스템 구축 및 현장 실증 (M4 transfer learning, M5 실공정 검증)'],
            ['Phase 3', '6개월', '통합 플랫폼 완성 및 기술 이전'],
        ]
    )
    
    doc.add_paragraph('')
    doc.add_heading('협력 체계', level=2)
    collab = [
        'Corning 천안: 실험 데이터 제공, ULE 7973 시편, 현장 검증 지원',
        'SPMDL: 시뮬레이터 개발, ML 모델링, 데이터 분석',
        'Advisory Board: 분기별 phase-gate 리뷰를 통한 연구 방향 관리',
    ]
    for c in collab:
        doc.add_paragraph(c, style='List Bullet')
    
    # =========================================================================
    # SECTION 6
    # =========================================================================
    doc.add_heading('6. 시뮬레이션 결과 미리보기', level=1)
    doc.add_paragraph(
        '아래 그림들은 현재 구현된 시뮬레이터의 대표적인 출력 결과입니다.'
    )
    
    figures = [
        ('fig1_nucleation_map.png', 'Figure 1. Nucleation Probability Map (M1)',
         'ULE 7973 기판(152×152 mm) 위의 결함 분포와 nucleation 확률 heatmap. '
         'CTE 불균일성에 의한 thermoelastic stress와 Griffith criterion 기반 nucleation 확률을 보여줍니다.'),
        ('fig2_crack_growth.png', 'Figure 2. Subcritical Crack Growth (M2)',
         '(a) 시간에 따른 균열 길이 성장 궤적. (b) Charles-Hillig V–K_I diagram으로 '
         'subcritical 영역(K₀ < K_I < K_IC)에서의 crack velocity를 보여줍니다.'),
        ('fig3_inspection_comparison.png', 'Figure 3. Inspection Method Comparison (M3)',
         '6종 비파괴 검사 기법(Acoustic, Laser Scattering, Raman, 193nm Interferometry, '
         'Electron Energy Loss Spectroscopy, Kelvin Force Microscopy)의 비교.'),
        ('fig4_ml_diagnostics.png', 'Figure 4. Bayesian ML Diagnostics (M4)',
         '(a) Bayesian 역추론 모델의 예측 정확도. (b) Physics-informed feature importance 분석.'),
        ('fig5_attribution.png', 'Figure 5. Process Attribution (M5)',
         '(a) Overlay 분산 분해: Scanner, Mask (Pristine), Mask (Degradation), Process 기여도. '
         '(b) Bayesian changepoint detection에 의한 degradation onset 자동 감지.'),
        ('fig6_material_comparison.png', 'Figure 6. Material Comparison',
         'ULE 7973, Zerodur, Clearceram-Z, AGC AZ, Shin-Etsu Quartz 5종 소재의 '
         'CTE 균일성, EUV 호환성, 치수 안정성, SCG exponent, 파괴 인성 비교 레이더 차트.'),
    ]
    
    for fname, caption, desc in figures:
        fpath = os.path.join(FIGDIR, fname)
        doc.add_paragraph(desc)
        if os.path.exists(fpath):
            doc.add_picture(fpath, width=Cm(15))
            cap_p = doc.add_paragraph(caption)
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_p.runs[0].italic = True
            cap_p.runs[0].font.size = Pt(10)
        doc.add_paragraph('')
    
    # =========================================================================
    # FOOTER
    # =========================================================================
    doc.add_paragraph('')
    footer = doc.add_paragraph(
        '본 제안서의 시뮬레이션 결과는 공개 문헌 기반의 물성치를 사용하였으며, '
        'Corning 독점 데이터는 포함되어 있지 않습니다. 실제 프로젝트 수행 시 '
        'Corning으로부터 제공받는 실험 데이터를 통해 모델 calibration을 진행할 예정입니다.'
    )
    footer.runs[0].italic = True
    footer.runs[0].font.size = Pt(9)
    
    # Save
    doc.save(OUTPATH)
    print(f"✓ DOCX saved: {OUTPATH}")


if __name__ == '__main__':
    create_docx()
